import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from models import Block

CCC_BLUE = RGBColor(0x00, 0xAE, 0xEF)
GRAY = RGBColor(0x73, 0x72, 0x72)
BODY = RGBColor(0x22, 0x22, 0x22)


def slug(text):
    s = re.sub(r"[^A-Za-z0-9]+", "-", text or "lesson").strip("-")
    return s.lower() or "lesson"


def _bottom_border(paragraph, color="00AEEF", size=18):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _run(paragraph, text, size=10.5, bold=False, color=BODY,
         caps=False):
    r = paragraph.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.all_caps = caps
    return r


def build_docx(doc, lang="en"):
    """Return a BytesIO containing the rendered lesson."""
    blocks = (Block.query.filter_by(source_doc_id=doc.id)
              .order_by(Block.seq).all())

    d = Document()

    for s in d.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    normal = d.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY

    title = d.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    if lang == "es":
        heading_text = doc.title_es or doc.title or doc.filename
    else:
        heading_text = doc.title or doc.filename
    _run(title, heading_text, size=20, bold=True, color=CCC_BLUE)

    rule = d.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(14)
    _bottom_border(rule)

    for b in blocks:
        text = (b.text_es if lang == "es" else b.text_en) or ""
        text = text.strip()
        if not text:
            continue

        if b.block_type == "heading":
            p = d.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            _run(p, text, size=11, bold=True, caps=True)

        elif b.block_type == "list_item":
            p = d.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            _run(p, text)

        else:
            p = d.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            _run(p, text)

    foot = d.sections[0].footer.paragraphs[0]
    if lang == "es":
        foot_text = "Revisión %s  |  Español (México)" % (
            doc.revision or "1.0")
    else:
        foot_text = "Revision %s  |  English" % (doc.revision or "1.0")
    _run(foot, foot_text, size=8, color=GRAY)

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf


def docx_filename(doc, lang="en"):
    return "%s-%s.docx" % (slug(doc.title or doc.filename), lang.upper())
